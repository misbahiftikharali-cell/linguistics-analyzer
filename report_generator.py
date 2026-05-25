import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

def safe_get_list(data, key):
    """Safely gets a list from a dict or returns the data if it is already a list."""
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        return data.get(key, [])
    return []

def add_stat_row(table, label, value):
    """Adds a bold label and value row to a table."""
    cells = table.add_row().cells
    cells[0].text = str(label)
    cells[1].text = str(value)
    cells[0].paragraphs[0].runs[0].bold = True

def generate_docx_report(results):
    """
    Generates a comprehensive, agent-by-agent forensic linguistic report.
    """
    doc = Document()
    
    # Global Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Title Page
    title = doc.add_heading('PMDD Forensic Linguistic Investigation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Multi-Agent Pragmatic Meaning Drift Detection Pipeline Output").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ==========================================
    # AGENT 5: EXECUTIVE SUMMARY (SYNTHESIS FIRST)
    # ==========================================
    doc.add_heading('Phase 0: Executive Synthesis & Verdict (Agent 5)', level=1)
    a5 = results.get("Agent5", {})
    if isinstance(a5, dict):
        exec_sum = a5.get("executive_summary", {})
        if exec_sum:
            doc.add_heading('Executive Overview', level=2)
            doc.add_paragraph(exec_sum.get("overview", "No overview available."))
            
            doc.add_heading('Major Findings', level=3)
            for finding in exec_sum.get("major_findings", []):
                doc.add_paragraph(finding, style='List Bullet')
            
            doc.add_heading('Dominant Patterns', level=3)
            for pattern in exec_sum.get("dominant_patterns", []):
                doc.add_paragraph(pattern, style='List Bullet')

        # Meaning Drift Scores
        doc.add_heading('Meaning Drift Assessment', level=2)
        scores = a5.get("drift_scores", {})
        if isinstance(scores, dict):
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Dimension'
            hdr[1].text = 'Score'
            hdr[2].text = 'Linguistic Reasoning'
            for cell in hdr: cell.paragraphs[0].runs[0].bold = True
            
            for key in ['pragmatic_drift', 'semantic_drift', 'register_drift', 'statistical_reliability']:
                data = scores.get(key, {})
                row = table.add_row().cells
                row[0].text = key.replace('_', ' ').title()
                row[1].text = f"{data.get('score', 0)}/100"
                row[2].text = data.get('reasoning', 'N/A')

        overall = scores.get("overall_meaning_drift", {})
        doc.add_heading(f"FINAL VERDICT: {overall.get('classification', 'N/A')} (Score: {overall.get('score', 0)})", level=2)
        doc.add_paragraph(overall.get('reasoning', ''), style='Intense Quote')

    doc.add_page_break()

    # ==========================================
    # AGENT 1: PREPROCESSING & SEGMENTATION
    # ==========================================
    doc.add_heading('Phase 1: Corpus Preprocessing (Agent 1)', level=1)
    a1_data = results.get("Agent1", {})
    a1_segments = safe_get_list(a1_data, "segments")
    info = a1_data.get("corpus_info", {}) if isinstance(a1_data, dict) else {}
        
    doc.add_paragraph(f"Total Segments Identified: {len(a1_segments)}")
    if info:
        doc.add_paragraph(f"Total Tokens: {info.get('total_tokens', 'N/A')} | Total Sentences: {info.get('total_sentences', 'N/A')}")
    
    if a1_segments:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'ID'
        hdr[1].text = 'Text Preview'
        hdr[2].text = 'Pos'
        hdr[3].text = 'Words'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        
        for seg in a1_segments:
            row = table.add_row().cells
            row[0].text = str(seg.get('id', 'N/A'))
            row[1].text = (seg.get('text', '')[:100] + '...') if len(seg.get('text', '')) > 100 else seg.get('text', '')
            row[2].text = str(seg.get('position', 'N/A'))
            row[3].text = str(seg.get('word_count', 'N/A'))
    
    doc.add_page_break()

    # ==========================================
    # AGENT 2: PRAGMATIC ANALYZER
    # ==========================================
    doc.add_heading('Phase 2: Pragmatic Evidence Map (Agent 2)', level=1)
    a2_results = safe_get_list(results.get("Agent2", []), "segments")
    
    if a2_results:
        # --- THEORY 1: Speech Act Theory ---
        doc.add_heading('Theory 1: Speech Act Theory (Austin/Searle)', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'ID'
        hdr[1].text = 'Speech Act'
        hdr[2].text = 'Forensic Explanation'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        
        for seg in a2_results:
            basis = f"Basis: {', '.join(seg.get('trigger_phrases', []))}" if seg.get('trigger_phrases') else ""
            row = table.add_row().cells
            row[0].text = str(seg.get('segment_id', 'N/A'))
            row[1].text = seg.get('speech_act', 'N/A')
            row[2].text = f"[{basis}] {seg.get('speech_act_explanation', seg.get('explanation', 'N/A'))}"
        
        if a2_results[0].get('speech_act_rules'):
            doc.add_heading('Speech Act Rules & Application', level=3)
            doc.add_paragraph(f"Rules: {a2_results[0].get('speech_act_rules')}")
            doc.add_paragraph(f"Application: {a2_results[0].get('speech_act_application')}")

        # --- THEORY 2: Gricean Maxims ---
        doc.add_heading('Theory 2: Cooperative Principle (Grice 1975)', level=2)
        grice_results = [s for s in a2_results if s.get('maxim_violations')]
        if grice_results:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'ID'
            hdr[1].text = 'Violation'
            hdr[2].text = 'Forensic Explanation'
            for cell in hdr: cell.paragraphs[0].runs[0].bold = True
            for seg in grice_results:
                row = table.add_row().cells
                row[0].text = str(seg.get('segment_id', 'N/A'))
                row[1].text = ", ".join(seg.get('maxim_violations', []))
                row[2].text = seg.get('maxim_explanation', 'N/A')
            
            if a2_results[0].get('maxim_rules'):
                doc.add_heading('Gricean Rules & Application', level=3)
                doc.add_paragraph(f"Rules: {a2_results[0].get('maxim_rules')}")
                doc.add_paragraph(f"Application: {a2_results[0].get('maxim_application')}")
        else:
            doc.add_paragraph("No significant Gricean violations detected.")

        # --- THEORY 3: Politeness Theory ---
        doc.add_heading('Theory 3: Politeness & Face Theory (Brown & Levinson)', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'ID'
        hdr[1].text = 'Politeness'
        hdr[2].text = 'Forensic Explanation'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        for seg in a2_results:
            row = table.add_row().cells
            row[0].text = str(seg.get('segment_id', 'N/A'))
            row[1].text = f"{seg.get('politeness_score', 'N/A')}/5"
            row[2].text = seg.get('politeness_explanation', 'N/A')
        
        if a2_results[0].get('politeness_rules'):
            doc.add_heading('Politeness Rules & Application', level=3)
            doc.add_paragraph(f"Rules: {a2_results[0].get('politeness_rules')}")
            doc.add_paragraph(f"Application: {a2_results[0].get('politeness_application')}")

    doc.add_page_break()

    # ==========================================
    # AGENT 3: SEMANTIC & REGISTER DETECTOR
    # ==========================================
    doc.add_heading('Phase 3: Semantic & Register Analysis (Agent 3)', level=1)
    a3_results = safe_get_list(results.get("Agent3", []), "segments")
    
    if a3_results:
        # --- THEORY 1: Semantic Field Theory ---
        doc.add_heading('Theory 1: Semantic Field Theory (Lyons 1977)', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'ID'
        hdr[1].text = 'Dominant Fields'
        hdr[2].text = 'Forensic Explanation'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        
        for a3 in a3_results:
            sa = a3.get("semantic_analysis", {})
            row = table.add_row().cells
            row[0].text = str(a3.get('segment_id', 'N/A'))
            row[1].text = ", ".join(sa.get('dominant_fields', []))
            row[2].text = sa.get('field_reasoning', 'N/A')
        
        if a3_results[0].get('semantic_analysis', {}).get('theory_rules'):
            doc.add_heading('Semantic Field Rules & Application', level=3)
            doc.add_paragraph(f"Rules: {a3_results[0]['semantic_analysis'].get('theory_rules')}")
            doc.add_paragraph(f"Application: {a3_results[0]['semantic_analysis'].get('how_applied')}")

        # --- THEORY 2: Register Analysis ---
        doc.add_heading('Theory 2: Register Analysis (Halliday 1978)', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'ID'
        hdr[1].text = 'Register'
        hdr[2].text = 'Field/Tenor/Mode'
        hdr[3].text = 'Forensic Explanation'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        
        for a3 in a3_results:
            ra = a3.get("register_analysis", {})
            row = table.add_row().cells
            row[0].text = str(a3.get('segment_id', 'N/A'))
            row[1].text = ra.get('primary_register', 'N/A')
            row[2].text = f"{ra.get('field', 'N/A')} / {ra.get('tenor', 'N/A')} / {ra.get('mode', 'N/A')}"
            row[3].text = ra.get('reasoning', 'N/A')
        
        if a3_results[0].get('register_analysis', {}).get('theory_rules'):
            doc.add_heading('Register Analysis Rules & Application', level=3)
            doc.add_paragraph(f"Rules: {a3_results[0]['register_analysis'].get('theory_rules')}")
            doc.add_paragraph(f"Application: {a3_results[0]['register_analysis'].get('how_applied')}")
            
        # Lexical Traceability (consolidated)
        doc.add_heading('Lexical Traceability Map', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Word'
        hdr[1].text = 'Field'
        hdr[2].text = 'Contextual Meaning'
        hdr[3].text = 'Shift?'
        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
        
        seen_words = set()
        for a3 in a3_results:
            for item in a3.get("lexical_items", []):
                word = item.get('word', '')
                if word and word.lower() not in seen_words:
                    row = table.add_row().cells
                    row[0].text = word
                    row[1].text = item.get('dominant_field', '')
                    row[2].text = item.get('contextual_meaning', '')
                    row[3].text = 'YES' if item.get('semantic_shift_detected') else 'No'
                    seen_words.add(word.lower())

            # Discourse Profile Summary
            dp = a3.get("discourse_semantic_profile", {})
            if dp:
                doc.add_paragraph("\nIntegrated Discourse Profile:").bold = True
                doc.add_paragraph(dp.get("overall_interpretation", "N/A"), style='Intense Quote')
                if dp.get("semantic_tension"):
                    doc.add_paragraph(f"Detected Tensions: {', '.join(dp.get('semantic_tension', []))}")

    doc.add_page_break()

    # ==========================================
    # AGENT 4: CORPUS STATISTICIAN
    # ==========================================
    doc.add_heading('Phase 4: Quantitative Corpus Profiling (Agent 4)', level=1)
    a4 = results.get("Agent4", {})
    if isinstance(a4, dict):
        # 1. Stats Overview
        stats = a4.get("corpus_statistics", {})
        if isinstance(stats, dict):
            doc.add_heading('Corpus Overview (Sinclair Metrics)', level=2)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            add_stat_row(table, "Total Tokens", stats.get("total_tokens", 0))
            add_stat_row(table, "Total Types", stats.get("total_types", 0))
            add_stat_row(table, "Type-Token Ratio (TTR)", f"{stats.get('type_token_ratio', 0.0):.4f}")
            add_stat_row(table, "Lexical Density", f"{stats.get('lexical_density', 0.0):.2f}")
        
        # 2. Frequency Analysis
        doc.add_heading('Frequency Analysis', level=2)
        freq_data = a4.get("frequency_analysis", [])
        if freq_data:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Word'
            hdr[1].text = 'Freq'
            hdr[2].text = 'Distribution Pattern'
            for cell in hdr: cell.paragraphs[0].runs[0].bold = True
            for item in freq_data[:15]: # Show top 15
                row = table.add_row().cells
                row[0].text = item.get("word", "")
                row[1].text = str(item.get("frequency", 0))
                row[2].text = item.get("distribution_pattern", "")
            
            if freq_data[0].get('theory_rules'):
                doc.add_heading('Frequency Analysis Rules & Application', level=3)
                doc.add_paragraph(f"Rules: {freq_data[0].get('theory_rules')}")
                doc.add_paragraph(f"Application: {freq_data[0].get('how_applied')}")

        # 3. Collocations
        doc.add_heading('Forensic Collocation Analysis', level=2)
        colloc_data = a4.get("collocation_analysis", [])
        if isinstance(colloc_data, list):
            for target in colloc_data:
                if not isinstance(target, dict): continue
                doc.add_paragraph(f"Target Word: {target.get('target_word')} (Freq: {target.get('corpus_frequency', 0)})").bold = True
                doc.add_paragraph(f"Drift Interpretation: {target.get('drift_interpretation', 'N/A')}", style='Intense Quote')
                
                collocs = target.get("collocates", [])
                if collocs and isinstance(collocs, list):
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    hdr = table.rows[0].cells
                    hdr[0].text = 'Collocate'
                    hdr[1].text = 'MI Score'
                    hdr[2].text = 'Strength'
                    for cell in hdr: cell.paragraphs[0].runs[0].bold = True
                    for ic in collocs[:5]: # Top 5
                        row = table.add_row().cells
                        row[0].text = str(ic.get('word', ''))
                        row[1].text = str(ic.get('mi_score', ''))
                        row[2].text = str(ic.get('association_strength', ''))

        # 4. Keyness & N-Grams
        doc.add_heading('Keyness & Formulaic Patterns', level=2)
        k_data = a4.get("keyness_analysis", [])
        n_data = a4.get("ngram_analysis", [])
        
        if k_data:
            doc.add_heading('Keyness (Scott Analysis)', level=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Word'
            hdr[1].text = 'Section'
            hdr[2].text = 'Interpretation'
            for cell in hdr: cell.paragraphs[0].runs[0].bold = True
            for k in k_data[:10]:
                row = table.add_row().cells
                row[0].text = k.get("word", "")
                row[1].text = k.get("dominant_section", "")
                row[2].text = k.get("keyness_interpretation", "")
            
            if k_data[0].get('theory_rules'):
                doc.add_heading('Keyness Analysis Rules & Application', level=3)
                doc.add_paragraph(f"Rules: {k_data[0].get('theory_rules')}")
                doc.add_paragraph(f"Application: {k_data[0].get('how_applied')}")

        if n_data:
            doc.add_heading('N-Gram Formulaic Expressions', level=3)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'N-Gram'
            hdr[1].text = 'Freq'
            hdr[2].text = 'Discourse Function'
            for cell in hdr: cell.paragraphs[0].runs[0].bold = True
            for n in n_data[:10]:
                row = table.add_row().cells
                row[0].text = n.get("ngram", "")
                row[1].text = str(n.get("frequency", 0))
                row[2].text = n.get("discourse_function", "")

    doc.add_page_break()

    # ==========================================
    # AGENT 5: DETAILED SYNTHESIS
    # ==========================================
    doc.add_heading('Phase 5: Integrated Forensic Synthesis (Agent 5)', level=1)
    if isinstance(a5, dict):
        # Integrated Findings
        layers = [
            ("Pragmatic Findings", a5.get("pragmatic_findings", {})),
            ("Semantic Findings", a5.get("semantic_findings", {})),
            ("Register Findings", a5.get("register_findings", {})),
            ("Quantitative Findings", a5.get("quantitative_findings", {}))
        ]
        
        for name, data in layers:
            if not data: continue
            doc.add_heading(name, level=2)
            doc.add_paragraph(data.get("interpretation", data.get("statistical_interpretation", "No interpretation available.")), style='Intense Quote')
            
            if data.get("theory_rules_summary"):
                doc.add_paragraph(f"Rules Applied: {data.get('theory_rules_summary')}")
                doc.add_paragraph(f"Application Logic: {data.get('application_logic_summary')}")

            # List some bullet points if they exist
            patterns = data.get("dominant_speech_acts", []) + data.get("semantic_shift_patterns", []) + data.get("register_borrowing_patterns", []) + data.get("collocational_patterns", [])
            for p in patterns:
                doc.add_paragraph(p, style='List Bullet')

        # Cross-Agent Synthesis
        doc.add_heading('Cross-Agent Synthesis', level=2)
        cas = a5.get("cross_agent_synthesis", {})
        if cas:
            doc.add_paragraph(cas.get("integrated_discourse_interpretation", "N/A"))
            doc.add_paragraph(f"Evidence Strength: {cas.get('evidence_strength', 'N/A')}")
            
        # Quality Control
        qc = a5.get("quality_control", {})
        if qc:
            doc.add_heading('Forensic Quality Control Report', level=2)
            doc.add_paragraph(f"Summary: {qc.get('quality_summary', 'N/A')}")
            if qc.get("contradictions_detected"):
                doc.add_paragraph(f"Contradictions: {', '.join(qc.get('contradictions_detected'))}")

        # Final Conclusion
        doc.add_heading('Conclusion & Recommendations', level=2)
        conc = a5.get("final_conclusion", {})
        doc.add_paragraph(conc.get("overall_interpretation", "N/A"), style='Normal')
        doc.add_paragraph(f"Linguistic Significance: {conc.get('linguistic_significance', 'N/A')}")
        doc.add_paragraph(f"Recommended Caution: {conc.get('recommended_caution', 'N/A')}")

    # Footer
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "PMDD Forensic Analysis System - Confidential Investigation Report"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to BytesIO
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
